"""Tests for the python-docx P0 parity round: Font (tri-state), ParagraphFormat,
TabStops, Cell/Table/Row/Section/Style extensions, Paragraph/Run methods.
"""
import io
import sys

import pytest

sys.path.insert(0, "tests")
from helpers import make_docx, read_docx_part, tp, p, run, cell, tr, tbl  # noqa: E402

from docxtplrs import DocxTemplate, Inches, Pt  # noqa: E402


def new_tpl(body=None, **kw):
    return DocxTemplate(io.BytesIO(make_docx(body or tp("x"), **kw)))


def saved_xml(tpl, part="word/document.xml"):
    out = io.BytesIO()
    tpl.save(out)
    return read_docx_part(out.getvalue(), part)


# ---------------- Font (tri-state) ----------------

def test_font_tri_state_semantics():
    tpl = new_tpl()
    r = tpl.paragraphs[0].add_run("x")
    f = r.font
    # missing -> None
    assert f.bold is None
    f.bold = True
    assert f.bold is True
    f.bold = False
    assert f.bold is False
    f.bold = None
    assert f.bold is None
    xml = saved_xml(tpl)
    assert "<w:b/>" not in xml and "<w:b " not in xml  # None removed the element


def test_font_tri_state_xml_forms():
    tpl = new_tpl()
    r = tpl.paragraphs[0].add_run("x")
    f = r.font
    f.bold = True
    f.italic = False
    xml = saved_xml(tpl)
    assert "<w:b/>" in xml  # True -> bare element
    assert '<w:i w:val="0"/>' in xml  # False -> val="0"


def test_font_all_tri_attrs():
    tpl = new_tpl()
    r = tpl.paragraphs[0].add_run("x")
    f = r.font
    attrs = [
        "bold", "cs_bold", "italic", "cs_italic", "all_caps", "small_caps",
        "strike", "double_strike", "outline", "shadow", "emboss", "imprint",
        "no_proof", "snap_to_grid", "hidden", "web_hidden", "spec_vanish",
        "rtl", "complex_script", "math",
    ]
    for a in attrs:
        assert getattr(f, a) is None, a
        setattr(f, a, True)
        assert getattr(f, a) is True, a
        setattr(f, a, None)
        assert getattr(f, a) is None, a


def test_font_size_name_color():
    tpl = new_tpl()
    r = tpl.paragraphs[0].add_run("x")
    f = r.font
    assert f.size is None
    f.size = Pt(14)
    assert f.size.emu == Pt(14).emu
    f.name = "Consolas"
    assert f.name == "Consolas"
    f.color.rgb = "#ff0000"
    assert f.color.rgb == "FF0000"
    f.color.rgb = None
    assert f.color.rgb is None
    xml = saved_xml(tpl)
    assert 'w:val="28"' in xml  # 14pt = 28 half-points
    assert 'w:ascii="Consolas"' in xml


def test_font_sub_superscript_exclusive():
    tpl = new_tpl()
    r = tpl.paragraphs[0].add_run("x")
    f = r.font
    f.subscript = True
    assert f.subscript is True and f.superscript is False
    f.superscript = True
    assert f.superscript is True and f.subscript is False
    f.superscript = None
    assert f.superscript is None


def test_font_underline_and_highlight():
    tpl = new_tpl()
    r = tpl.paragraphs[0].add_run("x")
    f = r.font
    f.underline = True
    assert f.underline == 1  # single
    f.underline = 3  # double
    assert f.underline == 3
    xml = saved_xml(tpl)
    assert '<w:u w:val="double"/>' in xml
    f.underline = None
    assert f.underline is None
    f.highlight_color = 7  # yellow
    assert f.highlight_color == 7
    xml = saved_xml(tpl)
    assert '<w:highlight w:val="yellow"/>' in xml


def test_style_font_and_paragraph_format():
    styles = '<w:style w:type="paragraph" w:styleId="MyP"><w:name w:val="My P"/></w:style>'
    tpl = new_tpl(styles=styles)
    st = tpl.styles["My P"]
    st.font.bold = True
    st.font.size = Pt(18)
    st.paragraph_format.alignment = 1  # center
    st.paragraph_format.space_before = Pt(12)
    xml = saved_xml(tpl, "word/styles.xml")
    assert "<w:b/>" in xml and 'w:val="36"' in xml
    assert '<w:jc w:val="center"/>' in xml and 'w:before="240"' in xml


# ---------------- ParagraphFormat ----------------

def test_paragraph_format_indents_spacing():
    tpl = new_tpl()
    pf = tpl.paragraphs[0].paragraph_format
    pf.left_indent = Inches(1)
    pf.right_indent = Inches(0.5)
    assert pf.left_indent.emu == Inches(1).emu
    assert abs(pf.right_indent.emu - Inches(0.5).emu) < 1000
    pf.first_line_indent = Inches(0.25)
    assert abs(pf.first_line_indent.emu - Inches(0.25).emu) < 1000
    pf.first_line_indent = Inches(-0.25)  # negative -> hanging
    assert abs(pf.first_line_indent.emu + Inches(0.25).emu) < 1000
    pf.space_before = Pt(6)
    pf.space_after = Pt(12)
    xml = saved_xml(tpl)
    assert 'w:left="1440"' in xml and 'w:right="720"' in xml
    assert 'w:hanging="360"' in xml
    assert 'w:before="120"' in xml and 'w:after="240"' in xml


def test_paragraph_format_alignment_and_flags():
    tpl = new_tpl()
    p0 = tpl.paragraphs[0]
    assert p0.alignment is None
    p0.alignment = 1
    assert p0.alignment == 1
    p0.alignment = "both"  # xml name accepted
    assert p0.alignment == 3
    pf = p0.paragraph_format
    assert pf.keep_together is None
    pf.keep_together = True
    pf.keep_with_next = True
    pf.page_break_before = True
    # widow_control defaults to True when missing (python-docx)
    assert pf.widow_control is True
    pf.widow_control = False
    assert pf.widow_control is False
    xml = saved_xml(tpl)
    for frag in ("<w:keepLines/>", "<w:keepNext/>", "<w:pageBreakBefore/>", '<w:widowControl w:val="0"/>'):
        assert frag in xml, frag


def test_paragraph_format_line_spacing():
    tpl = new_tpl()
    pf = tpl.paragraphs[0].paragraph_format
    assert pf.line_spacing is None
    pf.line_spacing = 2.0  # double
    assert abs(pf.line_spacing - 2.0) < 0.01
    assert pf.line_spacing_rule == 2  # DOUBLE
    pf.line_spacing = Pt(20)  # Length -> exact
    assert pf.line_spacing.emu == Pt(20).emu
    assert pf.line_spacing_rule == 4  # EXACTLY
    pf.line_spacing_rule = 3  # AT_LEAST keeps the length
    assert pf.line_spacing.emu == Pt(20).emu
    pf.line_spacing = None
    assert pf.line_spacing is None


def test_tab_stops():
    tpl = new_tpl()
    ts = tpl.paragraphs[0].paragraph_format.tab_stops
    assert len(ts) == 0
    t = ts.add_tab_stop(Inches(2), 1, 1)  # center, dots
    assert len(ts) == 1
    assert t.position.emu == Inches(2).emu
    assert t.alignment == 1 and t.leader == 1
    xml = saved_xml(tpl)
    assert '<w:tab w:val="center" w:leader="dot" w:pos="2880"/>' in xml
    ts.clear_all()
    assert len(ts) == 0


# ---------------- Paragraph / Run methods ----------------

def test_paragraph_clear_and_insert_before():
    tpl = new_tpl(tp("first") + tp("second"))
    p1 = tpl.paragraphs[1]
    new_p = p1.insert_paragraph_before("inserted")
    assert [p.text for p in tpl.paragraphs] == ["first", "inserted", "second"]
    assert new_p.text == "inserted"
    tpl.paragraphs[0].clear()
    assert tpl.paragraphs[0].text == ""
    xml = saved_xml(tpl)
    assert "<w:p>" in xml  # cleared paragraph kept (with no runs)


def test_run_add_break_tab_text_clear():
    tpl = new_tpl()
    r = tpl.paragraphs[0].add_run("a")
    r.add_break()       # line
    r.add_break(7)      # page
    r.add_tab()
    r.add_text("b")
    xml = saved_xml(tpl)
    assert "<w:br/>" in xml and '<w:br w:type="page"/>' in xml and "<w:tab/>" in xml
    assert r.text == "a\n\n\tb"  # w:br reads back as \n, w:tab as \t
    r.clear()
    assert r.text == ""


def test_run_text_setter_expands():
    tpl = new_tpl()
    r = tpl.paragraphs[0].add_run("x")
    r.text = "a\tb\nc"
    xml = saved_xml(tpl)
    assert "<w:tab/>" in xml and "<w:br/>" in xml
    assert r.text == "a\tb\nc"


# ---------------- Cell / Table / Row ----------------

def test_cell_paragraphs_and_add():
    tpl = new_tpl(tbl([tr(cell(tp("a")), cell(tp("b"))), tr(cell(tp("c")), cell(tp("d")))]))
    c = tpl.tables[0].cell(0, 0)
    assert [p.text for p in c.paragraphs] == ["a"]
    p2 = c.add_paragraph("second")
    assert [p.text for p in c.paragraphs] == ["a", "second"]
    p2.text = "changed"
    assert c.paragraphs[1].text == "changed"
    assert c.text == "achanged"


def test_cell_width_and_valign():
    tpl = new_tpl(tbl([tr(cell(tp("a")), cell(tp("b")))]))
    c = tpl.tables[0].cell(0, 0)
    c.width = Inches(2)
    assert c.width.emu == Inches(2).emu
    c.vertical_alignment = 1  # center
    assert c.vertical_alignment == 1
    xml = saved_xml(tpl)
    assert 'w:type="dxa"' in xml and 'w:w="2880"' in xml
    assert '<w:vAlign w:val="center"/>' in xml


def test_table_alignment_autofit_columns():
    tpl = new_tpl(tbl([tr(cell(tp("a")), cell(tp("b"))), tr(cell(tp("c")), cell(tp("d")))]))
    t = tpl.tables[0]
    assert t.autofit is True
    t.autofit = False
    assert t.autofit is False
    t.alignment = 1
    assert t.alignment == 1
    assert len(t.columns) == 2
    t.columns[0].width = Inches(1)
    assert t.columns[0].width.emu == Inches(1).emu
    col = t.add_column(Inches(3))
    assert len(t.columns) == 3
    assert col.width.emu == Inches(3).emu
    assert len(t.rows[0].cells) == 3
    assert [c.text for c in t.column_cells(0)] == ["a", "c"]
    assert [c.text for c in t.row_cells(1)] == ["c", "d", ""]
    xml = saved_xml(tpl)
    assert '<w:jc w:val="center"/>' in xml
    assert '<w:tblLayout w:type="fixed"/>' in xml


def test_row_height_and_rule():
    tpl = new_tpl(tbl([tr(cell(tp("a")), cell(tp("b")))]))
    row = tpl.tables[0].rows[0]
    assert row.height is None and row.height_rule is None
    row.height = Inches(1)
    row.height_rule = 2  # exact
    assert row.height.emu == Inches(1).emu
    assert row.height_rule == 2
    xml = saved_xml(tpl)
    assert '<w:trHeight w:val="1440" w:hRule="exact"/>' in xml


# ---------------- Section / Style ----------------

def test_section_start_type_and_distances():
    tpl = new_tpl()
    s = tpl.sections[0]
    assert s.start_type == 2  # missing w:type -> nextPage
    s.start_type = 0  # continuous
    assert s.start_type == 0
    s.start_type = "oddPage"
    assert s.start_type == 4
    s.start_type = None  # back to default
    assert s.start_type == 2
    s.header_distance = Inches(0.5)
    s.footer_distance = Inches(0.6)
    s.gutter = Inches(0.2)
    assert abs(s.header_distance.emu - Inches(0.5).emu) < 1000
    assert abs(s.footer_distance.emu - Inches(0.6).emu) < 1000
    assert abs(s.gutter.emu - Inches(0.2).emu) < 1000
    xml = saved_xml(tpl)
    assert 'w:header="720"' in xml and 'w:footer="864"' in xml and 'w:gutter="288"' in xml


def test_style_flags_and_next():
    styles = (
        '<w:style w:type="paragraph" w:styleId="A"><w:name w:val="A"/></w:style>'
        '<w:style w:type="paragraph" w:customStyle="1" w:styleId="B"><w:name w:val="B"/></w:style>'
    )
    tpl = new_tpl(styles=styles)
    a = tpl.styles["A"]
    b = tpl.styles["B"]
    assert a.builtin is True
    assert b.builtin is False
    assert a.hidden is False
    a.hidden = True
    assert a.hidden is True
    a.locked = True
    a.quick_style = True
    a.unhide_when_used = True
    a.priority = 10
    a.next_paragraph_style = "B"
    assert a.next_paragraph_style == "B"
    xml = saved_xml(tpl, "word/styles.xml")
    for frag in ("<w:semiHidden/>", "<w:locked/>", "<w:qFormat/>", "<w:unhideWhenUsed/>",
                 '<w:uiPriority w:val="10"/>', '<w:next w:val="B"/>'):
        assert frag in xml, frag
    a.next_paragraph_style = None
    assert a.next_paragraph_style is None


# ---------------- python-docx crosscheck ----------------

def test_python_docx_reads_our_output():
    """Properties written through docxtplrs must be readable by python-docx."""
    docx = pytest.importorskip("docx")
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt as DPt

    tpl = new_tpl(tp("hello"))
    p0 = tpl.paragraphs[0]
    p0.alignment = 1
    p0.paragraph_format.space_after = Pt(18)
    p0.paragraph_format.line_spacing = 1.5
    r = p0.add_run("world")
    r.font.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = "00FF00"
    out = io.BytesIO()
    tpl.save(out)
    out.seek(0)

    doc = docx.Document(out)
    p = doc.paragraphs[0]
    assert p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert p.paragraph_format.space_after == DPt(18)
    assert p.paragraph_format.line_spacing == 1.5
    run = p.runs[-1]
    assert run.font.bold is True
    assert run.font.size == DPt(16)
    assert str(run.font.color.rgb) == "00FF00"
